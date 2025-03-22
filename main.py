#!/usr/bin/env python3
"""
Resume Builder Web Application using Groq API
This web application allows users to input their resume and job description
and generates a customized resume tailored to the job posting using the Groq API.
"""

import os
import datetime
import logging
import json
import magic
import traceback
import hashlib
import re
from typing import Optional, Tuple, List, Dict
from pathlib import Path
from flask import Flask, render_template, request, jsonify, flash, redirect, url_for, send_file
from groq import Groq
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.utils import secure_filename

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-here')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  
ALLOWED_EXTENSIONS = {'pdf'}

MODEL_NAME = "llama-3.3-70b-versatile"
OUTPUT_DIR = "output"
MAX_TOKENS = 1024

GREEN_JOB_CATEGORIES = {
    "renewable_energy": ["solar", "wind", "hydro", "geothermal", "renewable", "clean energy", "solar panel", "wind turbine", "energy storage"],
    "sustainability": ["sustainable", "environmental", "eco-friendly", "green", "carbon footprint", "net zero", "sustainability", "environmental impact"],
    "conservation": ["conservation", "wildlife", "biodiversity", "ecosystem", "habitat", "endangered species", "environmental protection", "natural resources"],
    "clean_tech": ["clean technology", "green technology", "smart grid", "energy efficiency", "electric vehicles", "battery technology", "smart home", "green building"],
    "waste_management": ["recycling", "waste reduction", "circular economy", "zero waste", "composting", "wastewater", "landfill", "waste diversion"],
    "climate_action": ["climate change", "carbon reduction", "emissions", "climate policy", "carbon trading", "climate adaptation", "climate mitigation"],
    "green_transportation": ["electric vehicles", "public transit", "bicycle infrastructure", "sustainable transport", "green logistics", "fleet electrification"],
    "sustainable_agriculture": ["organic farming", "permaculture", "sustainable food", "local food", "food security", "regenerative agriculture", "vertical farming"]
}

ATS_KEYWORDS = {
    "software": ["python", "javascript", "java", "sql", "agile", "scrum", "git", "docker", "kubernetes", "aws", "azure"],
    "data": ["data analysis", "machine learning", "statistics", "sql", "python", "r", "tableau", "power bi", "excel"],
    "marketing": ["digital marketing", "seo", "social media", "content strategy", "analytics", "crm", "email marketing"],
    "finance": ["financial analysis", "excel", "accounting", "risk management", "financial modeling", "forecasting"],
    "healthcare": ["patient care", "medical records", "healthcare management", "clinical", "healthcare compliance"],
    "education": ["curriculum development", "teaching", "student engagement", "lesson planning", "assessment"],
    "sales": ["sales strategy", "client relationship", "negotiation", "crm", "sales pipeline", "revenue growth"],
    "hr": ["recruitment", "talent management", "employee relations", "hr policies", "onboarding", "training"],
    "operations": ["process improvement", "project management", "supply chain", "logistics", "quality control"],
    "customer_service": ["customer support", "client satisfaction", "problem resolution", "communication", "service delivery"]
}

def setup_environment() -> None:
    """Load environment variables and create necessary directories."""
    try:
        load_dotenv()
        if not os.getenv('GROQ_API_KEY'):
            raise ValueError("GROQ_API_KEY environment variable is not set")
        
        Path(OUTPUT_DIR).mkdir(exist_ok=True)
        Path(app.config['UPLOAD_FOLDER']).mkdir(exist_ok=True)
        logger.info("Environment setup completed successfully")
    except Exception as e:
        logger.error(f"Environment setup failed: {str(e)}")
        raise

def allowed_file(filename: str) -> bool:
    """Check if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF
    """
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        logger.info(f"Successfully extracted text from PDF: {pdf_path}")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise

def create_pdf_from_text(text: str, output_path: str, template: str = 'modern') -> None:
    """
    Create a PDF file from text content.
    
    Args:
        text (str): Text content to write to PDF
        output_path (str): Path where to save the PDF
        template (str): Template style to use
    """
    try:
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        if template == 'modern':
            title_font_size = 24
            heading_font_size = 16
            body_font_size = 12
        elif template == 'classic':
            title_font_size = 20
            heading_font_size = 14
            body_font_size = 12
        else:  
            title_font_size = 28
            heading_font_size = 18
            body_font_size = 12
        
        c.setFont("Helvetica-Bold", title_font_size)
        title = "Professional Resume"
        c.drawString(width/2 - c.stringWidth(title, "Helvetica-Bold", title_font_size)/2, height - 50, title)
        
        c.setFont("Helvetica", body_font_size)
        
        lines = text.split('\n')
        y = height - 100 
        
        for line in lines:
            if y < 50: 
                c.showPage()
                c.setFont("Helvetica", body_font_size)
                y = height - 50
            
            c.drawString(50, y, line)
            y -= 15 
        
        c.save()
        logger.info(f"Successfully created PDF at: {output_path}")
    except Exception as e:
        logger.error(f"Error creating PDF: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise

def create_docx_from_text(text: str, output_path: str, template: str = 'modern') -> None:
    """
    Create a DOCX file from text content.
    
    Args:
        text (str): Text content to write to DOCX
        output_path (str): Path where to save the DOCX
        template (str): Template style to use
    """
    try:
        doc = Document()
        
        title = doc.add_heading('Professional Resume', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for line in text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                if template == 'modern':
                    p.style = 'Normal'
                elif template == 'classic':
                    p.style = 'Body Text'
                else:  
                    p.style = 'Normal'
        
        doc.save(output_path)
        logger.info(f"Successfully created DOCX at: {output_path}")
    except Exception as e:
        logger.error(f"Error creating DOCX: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise

def extract_keywords(text: str) -> List[str]:
    """
    Extract important keywords from text.
    
    Args:
        text (str): Text to analyze
        
    Returns:
        List[str]: List of extracted keywords
    """
    try:
        
        words = text.lower().split()
        
        common_words = {'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'i', 'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you', 'do', 'at'}
        words = [word.strip('.,!?()[]{}":;') for word in words if word not in common_words]
        
        word_freq = {}
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, _ in keywords[:20]]  
    except Exception as e:
        logger.error(f"Error extracting keywords: {str(e)}")
        return []

def match_job_opportunities(resume_keywords: List[str], job_description: str) -> List[Dict[str, str]]:
    """
    Match resume keywords with job opportunities.
    
    Args:
        resume_keywords (List[str]): Keywords from the resume
        job_description (str): The job description text
        
    Returns:
        List[Dict[str, str]]: List of matching job opportunities
    """
    try:
        matches = []
        job_desc_lower = job_description.lower()
        resume_keywords_lower = [k.lower() for k in resume_keywords]
        
        for industry, keywords in ATS_KEYWORDS.items():
            matching_keywords = [keyword for keyword in keywords if keyword in job_desc_lower]
            resume_matches = [keyword for keyword in keywords if keyword in resume_keywords_lower]
            
            if matching_keywords or resume_matches:
                
                job_match_percentage = len(matching_keywords) / len(keywords) * 100 if matching_keywords else 0
                resume_match_percentage = len(resume_matches) / len(keywords) * 100 if resume_matches else 0
                overall_match = (job_match_percentage + resume_match_percentage) / 2
                
                recommendation = f"This position has a {round(overall_match, 2)}% overall match with your experience in {industry.replace('_', ' ')}. "
                
                if job_match_percentage > 0:
                    recommendation += f"\n- Job requirements match: {round(job_match_percentage, 2)}% "
                    recommendation += f"(matching keywords: {', '.join(matching_keywords)})"
                
                if resume_match_percentage > 0:
                    recommendation += f"\n- Your experience match: {round(resume_match_percentage, 2)}% "
                    recommendation += f"(matching keywords: {', '.join(resume_matches)})"
                
                if overall_match > 70:
                    recommendation += "\n- Strong match: Your experience aligns well with the job requirements."
                elif overall_match > 40:
                    recommendation += "\n- Moderate match: Consider highlighting relevant experience and skills."
                else:
                    recommendation += "\n- Limited match: Focus on transferable skills and relevant achievements."
                
                matches.append({
                    "industry": industry.replace("_", " ").title(),
                    "overall_match_percentage": round(overall_match, 2),
                    "job_match_percentage": round(job_match_percentage, 2),
                    "resume_match_percentage": round(resume_match_percentage, 2),
                    "matching_keywords": matching_keywords,
                    "resume_matching_keywords": resume_matches,
                    "recommendation": recommendation
                })
        
        matches.sort(key=lambda x: x["overall_match_percentage"], reverse=True)
        return matches
    except Exception as e:
        logger.error(f"Error matching job opportunities: {str(e)}")
        return []

def optimize_for_ats(resume: str, job_description: str) -> str:
    """
    Optimize resume content for ATS systems.
    
    Args:
        resume (str): Original resume text
        job_description (str): Job description text
        
    Returns:
        str: Optimized resume text
    """
    try:
        client = Groq(api_key=os.environ.get('GROQ_API_KEY'))
        
        job_keywords = extract_keywords(job_description)
        
        prompt = f"""
        Optimize this resume for ATS systems and improve grammar. The job description keywords are: {', '.join(job_keywords)}
        
        Original Resume:
        {resume}
        
        Please:
        1. Incorporate relevant keywords naturally and contextually
        2. Improve grammar, punctuation, and sentence structure
        3. Use strong action verbs and quantifiable achievements
        4. Maintain professional tone and clarity
        5. Keep the same information but make it more impactful
        6. Ensure proper formatting and section organization
        7. Add industry-specific terminology where appropriate
        8. Remove any jargon or unclear language
        9. Make achievements more specific and measurable
        10. Ensure consistency in tense and voice throughout
        
        Format the output with clear sections:
        - Professional Summary
        - Work Experience
        - Education
        - Skills
        - Additional Qualifications
        """
        
        completion = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": "I'll help optimize your resume for ATS systems and improve its overall quality."}
            ],
            temperature=0.7,
            max_tokens=MAX_TOKENS,
            top_p=1,
            stream=False
        )
        
        optimized_content = completion.choices[0].message.content
        logger.info("Successfully optimized resume content")
        return optimized_content
        
    except Exception as e:
        logger.error(f"Error optimizing resume: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return resume

def generate_custom_resume(resume: str, job_description: str) -> Optional[str]:
    """
    Generate a customized resume using the Groq API.
    
    Args:
        resume (str): The original resume text
        job_description (str): The job description text
        
    Returns:
        Optional[str]: The generated resume or None if there's an error
    """
    try:
        optimized_resume = optimize_for_ats(resume, job_description)
        
        client = Groq(api_key=os.environ.get('GROQ_API_KEY'))
        
        resume_keywords = extract_keywords(resume)
        job_matches = match_job_opportunities(resume_keywords, job_description)
        
        prompt = f"""
        Create a customized resume based on this optimized version and job description.
        
        Optimized Resume:
        {optimized_resume}
        
        Job Description:
        {job_description}
        
        Job Matches Analysis:
        {json.dumps(job_matches, indent=2)}
        
        Please:
        1. Tailor the content to match the job requirements
        2. Highlight relevant skills and experience based on the match analysis
        3. Use industry-specific terminology from the job description
        4. Maintain professional formatting and structure
        5. Keep the optimized ATS-friendly content
        6. Emphasize matching keywords naturally
        7. Add quantifiable achievements where possible
        8. Ensure all sections are properly organized
        9. Use action verbs and strong language
        10. Make the resume stand out while staying professional
        
        Format the output with clear sections:
        - Professional Summary (tailored to the job)
        - Work Experience (highlighting relevant achievements)
        - Education
        - Skills (prioritizing job-relevant skills)
        - Additional Qualifications
        """
        
        completion = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": "I'll help create a customized resume that matches the job requirements."}
            ],
            temperature=0.7,
            max_tokens=MAX_TOKENS,
            top_p=1,
            stream=False
        )
        
        generated_content = completion.choices[0].message.content
        logger.info("Successfully generated resume content")
        return generated_content
        
    except Exception as e:
        logger.error(f"Error generating resume: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return None

def save_output(content: str, template: str = 'modern') -> Tuple[Optional[str], Optional[str]]:
    """
    Save the generated content to files with timestamp.
    
    Args:
        content (str): Content to save
        template (str): Template style to use
        
    Returns:
        Tuple[Optional[str], Optional[str]]: Paths to the saved PDF and DOCX files
    """
    try:
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        
        pdf_file = Path(OUTPUT_DIR) / f"resume-{timestamp}.pdf"
        create_pdf_from_text(content, str(pdf_file), template)
        
        docx_file = Path(OUTPUT_DIR) / f"resume-{timestamp}.docx"
        create_docx_from_text(content, str(docx_file), template)
        
        logger.info(f"Output saved to: {pdf_file} and {docx_file}")
        return str(pdf_file), str(docx_file)
        
    except Exception as e:
        logger.error(f"Error saving output files: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return None, None

def analyze_green_job_opportunities(job_description: str) -> List[Dict[str, str]]:
    """
    Analyze job description for green job opportunities.
    
    Args:
        job_description (str): The job description text
        
    Returns:
        List[Dict[str, str]]: List of matching green job categories and recommendations
    """
    try:
        job_desc_lower = job_description.lower()
        matches = []
        
        for category, keywords in GREEN_JOB_CATEGORIES.items():
            matching_keywords = [keyword for keyword in keywords if keyword in job_desc_lower]
            if matching_keywords:
                match_percentage = len(matching_keywords) / len(keywords) * 100
                
                recommendation = f"This position has a {round(match_percentage, 2)}% match with {category.replace('_', ' ')} initiatives. "
                if match_percentage > 70:
                    recommendation += "This is a strong match for green job opportunities. "
                elif match_percentage > 40:
                    recommendation += "This position has significant green job potential. "
                else:
                    recommendation += "This position has some green job elements. "
                
                recommendation += "Consider highlighting relevant experience in this area."
                
                matches.append({
                    "category": category.replace("_", " ").title(),
                    "match_percentage": round(match_percentage, 2),
                    "keywords": matching_keywords,
                    "recommendation": recommendation
                })
        
        # Sort by match percentage
        matches.sort(key=lambda x: x["match_percentage"], reverse=True)
        return matches
    except Exception as e:
        logger.error(f"Error analyzing green job opportunities: {str(e)}")
        return []

def generate_blockchain_hash(content: str) -> Dict[str, str]:
    """
    Generate a blockchain hash for resume verification.
    
    Args:
        content (str): The resume content to hash
        
    Returns:
        Dict[str, str]: Dictionary containing hash and verification details
    """
    try:
        content_hash = hashlib.sha256(content.encode()).hexdigest()
        
        timestamp = datetime.datetime.now().isoformat()
        verification_data = {
            "content_hash": content_hash,
            "timestamp": timestamp,
            "version": "1.0",
            "verification_type": "resume",
            "hash_algorithm": "SHA-256"
        }
        
        verification_hash = hashlib.sha256(json.dumps(verification_data).encode()).hexdigest()
        
        return {
            "verification_hash": verification_hash,
            "timestamp": timestamp,
            "content_hash": content_hash,
            "verification_url": f"https://resume-verification.com/verify/{verification_hash}"
        }
    except Exception as e:
        logger.error(f"Error generating blockchain hash: {str(e)}")
        return {
            "verification_hash": "",
            "timestamp": "",
            "content_hash": "",
            "verification_url": ""
        }

@app.route('/build')
def build():
    """Render the main page."""
    return render_template('index.html')

@app.route('/')
def index():
    return render_template('home.html')

@app.route('/upload', methods=['POST'])
def upload_resume():
    """Handle resume file upload."""
    try:
        if 'resume' not in request.files:
            logger.error("No file part in request")
            return jsonify({'error': 'No file part'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            logger.error("No selected file")
            return jsonify({'error': 'No selected file'}), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            resume_text = extract_text_from_pdf(filepath)
            
            os.remove(filepath)
            
            return jsonify({
                'success': True,
                'resume_text': resume_text
            })
        
        logger.error(f"Invalid file type: {file.filename}")
        return jsonify({'error': 'Invalid file type'}), 400
        
    except Exception as e:
        logger.error(f"Error in upload_resume route: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/generate', methods=['POST'])
def generate_resume():
    """Handle resume generation request."""
    try:
        data = request.get_json()
        if not data:
            logger.error("No JSON data received")
            return jsonify({'error': 'No data received'}), 400
            
        if 'resume' not in data or 'job_description' not in data:
            logger.error(f"Missing required fields. Received data: {data}")
            return jsonify({'error': 'Missing required fields'}), 400
        
        resume = data['resume']
        job_description = data['job_description']
        template = data.get('template', 'modern')
        
        logger.info("Starting resume generation")
        
        resume_keywords = extract_keywords(resume)
        
        job_matches = match_job_opportunities(resume_keywords, job_description)
        
        generated_resume = generate_custom_resume(resume, job_description)
        
        if not generated_resume:
            logger.error("Failed to generate resume content")
            return jsonify({'error': 'Failed to generate resume'}), 500
        
        green_opportunities = analyze_green_job_opportunities(job_description)
        
        verification_data = generate_blockchain_hash(generated_resume)
        
        pdf_path, docx_path = save_output(generated_resume, template)
        
        if not pdf_path or not docx_path:
            logger.error("Failed to save output files")
            return jsonify({'error': 'Failed to save resume'}), 500
        
        return jsonify({
            'success': True,
            'resume': generated_resume,
            'pdf_path': pdf_path,
            'docx_path': docx_path,
            'green_opportunities': green_opportunities,
            'verification_data': verification_data,
            'job_matches': job_matches,
            'resume_keywords': resume_keywords
        })
        
    except Exception as e:
        logger.error(f"Error in generate_resume route: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/download/<path:filename>')
def download_file(filename):
    """Handle file download."""
    try:
        file_path = os.path.join(OUTPUT_DIR, filename)
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return jsonify({'error': 'File not found'}), 404
            
        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"Error downloading file: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': f'Error downloading file: {str(e)}'}), 500

@app.route('/suggest', methods=['POST'])
def get_suggestions():
    """Handle real-time AI suggestions request."""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
            
        text = data['text']
        section = data.get('section', 'general')  
        
        client = Groq(api_key=os.environ.get('GROQ_API_KEY'))
        
        section_prompts = {
            'experience': """
            Analyze this work experience section and provide suggestions for improvement:
            1. Use stronger action verbs
            2. Add quantifiable achievements
            3. Highlight relevant skills
            4. Improve clarity and impact
            5. Make it more ATS-friendly
            
            Current text:
            {text}
            
            Provide 3-5 specific suggestions for improvement.
            """,
            'education': """
            Analyze this education section and provide suggestions for improvement:
            1. Highlight relevant coursework
            2. Add academic achievements
            3. Include relevant certifications
            4. Make it more impactful
            5. Ensure proper formatting
            
            Current text:
            {text}
            
            Provide 3-5 specific suggestions for improvement.
            """,
            'skills': """
            Analyze this skills section and provide suggestions for improvement:
            1. Add relevant technical skills
            2. Include soft skills
            3. Prioritize most important skills
            4. Make it more specific
            5. Ensure ATS compatibility
            
            Current text:
            {text}
            
            Provide 3-5 specific suggestions for improvement.
            """,
            'general': """
            Analyze this resume section and provide suggestions for improvement:
            1. Improve clarity and impact
            2. Make it more professional
            3. Enhance readability
            4. Add relevant details
            5. Make it more ATS-friendly
            
            Current text:
            {text}
            
            Provide 3-5 specific suggestions for improvement.
            """
        }
        
        prompt = section_prompts.get(section, section_prompts['general']).format(text=text)
        
        completion = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": "I'll help improve your resume section with specific suggestions."}
            ],
            temperature=0.7,
            max_tokens=500,  
            top_p=1,
            stream=False
        )
        
        suggestions = completion.choices[0].message.content
        
        keywords = extract_keywords(text)
        
        return jsonify({
            'success': True,
            'suggestions': suggestions,
            'keywords': keywords
        })
        
    except Exception as e:
        logger.error(f"Error getting suggestions: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

@app.route('/create', methods=['GET', 'POST'])
def create_resume():
    if request.method == 'GET':
        return render_template('create.html')
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': 'No data received'
            }), 400
        
        app.logger.info(f"Received data: {data}")
        
        personal_info = {
            'first_name': data.get('firstName', ''),
            'last_name': data.get('lastName', ''),
            'email': data.get('email', ''),
            'phone': data.get('phone', ''),
            'location': data.get('location', '')
        }
        
        summary = data.get('summary', '')
        
        experience = []
        if not data.get('isFresher', False):
            companies = data.get('companies', [])
            positions = data.get('positions', [])
            startDates = data.get('startDates', [])
            endDates = data.get('endDates', [])
            responsibilities = data.get('responsibilities', [])
            
            for i in range(len(companies)):
                if companies[i] and positions[i] and startDates[i] and endDates[i] and responsibilities[i]:
                    experience.append({
                        'company': companies[i],
                        'position': positions[i],
                        'start_date': startDates[i],
                        'end_date': endDates[i],
                        'responsibilities': responsibilities[i]
                    })
        
        education = []
        institutions = data.get('institutions', [])
        degrees = data.get('degrees', [])
        eduStartDates = data.get('eduStartDates', [])
        eduEndDates = data.get('eduEndDates', [])
        
        for i in range(len(institutions)):
            if institutions[i] and degrees[i] and eduStartDates[i] and eduEndDates[i]:
                education.append({
                    'institution': institutions[i],
                    'degree': degrees[i],
                    'start_date': eduStartDates[i],
                    'end_date': eduEndDates[i]
                })
        
        skills = {
            'technical': data.get('technicalSkills', []),
            'soft': data.get('softSkills', [])
        }
        
        template = data.get('template', 'modern')
        
        resume_content = generate_resume_content(
            personal_info=personal_info,
            summary=summary,
            experience=experience,
            education=education,
            skills=skills,
            template=template
        )
        
        os.makedirs('output', exist_ok=True)
        
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        pdf_path = f"output/resume-{timestamp}.pdf"
        docx_path = f"output/resume-{timestamp}.docx"
        
        try:
            create_pdf_from_text(resume_content, pdf_path)
            create_docx_from_text(resume_content, docx_path)
        except Exception as e:
            app.logger.error(f"Error generating files: {str(e)}")
            return jsonify({
                'success': False,
                'error': f"Error generating files: {str(e)}"
            }), 500
        
        return jsonify({
            'success': True,
            'resume': resume_content,
            'pdf_path': pdf_path,
            'docx_path': docx_path
        })
        
    except Exception as e:
        app.logger.error(f"Error creating resume: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

def generate_resume_content(personal_info, summary, experience, education, skills, template='modern'):
    """Generate resume content based on the provided information and template."""
    
    header = f"{personal_info.get('first_name', '')} {personal_info.get('last_name', '')}\n"
    header += f"{personal_info.get('email', '')} | {personal_info.get('phone', '')} | {personal_info.get('location', '')}\n\n"
    
    content = header
    content += "PROFESSIONAL SUMMARY\n"
    content += "=" * 50 + "\n"
    content += f"{summary}\n\n" if summary else "\n"
    
    if experience and any(exp.get('company') for exp in experience):
        content += "PROFESSIONAL EXPERIENCE\n"
        content += "=" * 50 + "\n"
        for exp in experience:
            if exp.get('company') and exp.get('position'):
                content += f"{exp.get('position', '')}\n"
                content += f"{exp.get('company', '')} | {exp.get('start_date', '')} - {exp.get('end_date', '')}\n"
                content += f"{exp.get('responsibilities', '')}\n\n"
    
    if education and any(edu.get('institution') for edu in education):
        content += "EDUCATION\n"
        content += "=" * 50 + "\n"
        for edu in education:
            if edu.get('institution') and edu.get('degree'):
                content += f"{edu.get('degree', '')}\n"
                content += f"{edu.get('institution', '')} | {edu.get('start_date', '')} - {edu.get('end_date', '')}\n\n"
    
    if skills.get('technical') or skills.get('soft'):
        content += "SKILLS\n"
        content += "=" * 50 + "\n"
        if skills.get('technical'):
            content += "Technical Skills: " + ", ".join(skills['technical']) + "\n"
        if skills.get('soft'):
            content += "Soft Skills: " + ", ".join(skills['soft']) + "\n"
    
    return content

if __name__ == "__main__":
    setup_environment()
    
    app.run(debug=True, port=5000)
